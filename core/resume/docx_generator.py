"""Generate a tailored 2-page DOCX resume using AI."""

import json
from pathlib import Path
from typing import Optional
from datetime import datetime
from logger import get_logger

logger = get_logger("resume.docx_generator")

TAILOR_PROMPT = """Tailor this resume for the target role. Never fabricate facts.

RESUME:
{resume_text}

JD ({company} — {role}):
{job_description}

MISSING KEYWORDS: {missing_keywords}

Rules: keep all facts accurate; reorder bullets by relevance; incorporate missing keywords naturally; rewrite summary for this role; use JD action verbs; keep existing section structure.

Return ONLY valid JSON:
{{"summary":"3-4 line summary","sections":[{{"heading":"Name","entries":[{{"title":"","subtitle":"","bullets":[]}}]}}],"skills":{{"category":["skill"]}}}}"""


def _extract_docx_text(docx_path: Path) -> str:
    """Extract plain text from DOCX preserving rough structure."""
    try:
        from docx import Document
        doc = Document(str(docx_path))
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        return "\n".join(lines)
    except Exception as e:
        logger.error("Failed to read DOCX %s: %s", docx_path, e)
        return ""


_DEFAULT_MODEL = "claude-haiku-4-5-20251001"  # Fast + cheap; override with TAILOR_MODEL env var

def _call_ai(prompt: str, model: str | None = None) -> Optional[str]:
    """Call Claude CLI via subprocess: `claude -p <prompt> --model <model>`."""
    import subprocess
    import shutil
    import re
    import time

    import os
    resolved_model = model or os.getenv("TAILOR_MODEL", _DEFAULT_MODEL)
    claude_bin = shutil.which("claude")
    if not claude_bin:
        logger.error("claude CLI not found in PATH — install Claude Code CLI")
        return None

    ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')

    for attempt in range(2):
        try:
            result = subprocess.run(
                [claude_bin, "-p", prompt, "--model", resolved_model],
                capture_output=True,
                text=True,
                timeout=120,
            )
            if result.returncode != 0:
                stderr = result.stderr or ""
                if "usage limit" in stderr.lower():
                    logger.error("Claude usage limit reached: %s", stderr[:200])
                    return None
                logger.error("claude CLI exited %d: %s", result.returncode, stderr[:300])
                if attempt < 1:
                    time.sleep(5)
                continue
            output = ansi_escape.sub('', result.stdout).strip()
            if not output:
                logger.error("claude CLI returned empty output. stderr: %s", result.stderr[:300])
                if attempt < 1:
                    time.sleep(5)
                continue
            return output
        except subprocess.TimeoutExpired:
            logger.error("claude CLI timed out after 120s (attempt %d/2)", attempt + 1)
            if attempt < 1:
                time.sleep(5)
        except Exception as e:
            logger.error("claude CLI subprocess error: %s", e)
            return None

    return None


def _build_docx(data: dict, candidate_name: str = "Tarun Gupta") -> Optional[object]:
    """Build a formatted DOCX Document from structured AI data."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        logger.error("python-docx not installed")
        return None

    doc = Document()

    # --- Page margins (narrow for 2-page fit) ---
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    def _set_paragraph_spacing(para, before=0, after=2):
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after = Pt(after)

    def _add_name_header():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, after=2)
        run = p.add_run(candidate_name)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    def _add_section_heading(text: str):
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=6, after=2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1A, 0x5F, 0xD4)
        # Underline via bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1A5FD4")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_entry(title: str, subtitle: str, bullets: list[str]):
        # Title + subtitle row
        if title or subtitle:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, after=1)
            if title:
                run = p.add_run(title)
                run.bold = True
                run.font.size = Pt(10)
            if subtitle:
                run2 = p.add_run(f"  |  {subtitle}" if title else subtitle)
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        for bullet in bullets:
            p = doc.add_paragraph(style="List Bullet")
            _set_paragraph_spacing(p, after=1)
            run = p.add_run(bullet)
            run.font.size = Pt(9)

    # Build document
    _add_name_header()

    # Summary
    summary = data.get("summary", "")
    if summary:
        _add_section_heading("Professional Summary")
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, after=3)
        run = p.add_run(summary)
        run.font.size = Pt(9)
        run.font.italic = True

    # Sections
    for section_data in data.get("sections", []):
        heading = section_data.get("heading", "")
        if not heading:
            continue
        _add_section_heading(heading)
        for entry in section_data.get("entries", []):
            _add_entry(
                entry.get("title", ""),
                entry.get("subtitle", ""),
                entry.get("bullets", []),
            )

    # Skills
    skills = data.get("skills", {})
    if skills:
        _add_section_heading("Skills")
        for category, skill_list in skills.items():
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, after=2)
            run_cat = p.add_run(f"{category}: ")
            run_cat.bold = True
            run_cat.font.size = Pt(9)
            run_skills = p.add_run(", ".join(skill_list))
            run_skills.font.size = Pt(9)

    return doc


def generate_tailored_resume(
    original_docx_path: Path,
    job: dict,
    output_dir: Path = Path("data/tailored_resumes"),
) -> Optional[Path]:
    """
    Generate a tailored 2-page DOCX resume for a job.

    Returns path to saved DOCX, or None on failure.
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    resume_text = _extract_docx_text(original_docx_path)
    if not resume_text:
        logger.error("Could not extract text from %s", original_docx_path)
        return None

    company = str(job.get("Company", job.get("company", "Unknown")))
    role = str(job.get("Role", job.get("title", "Unknown")))
    description = str(job.get("Description", job.get("description", "")))
    missing_skills = job.get("Missing Skills", job.get("missing_skills", ""))
    if isinstance(missing_skills, list):
        missing_keywords = ", ".join(missing_skills)
    else:
        missing_keywords = str(missing_skills)

    prompt = TAILOR_PROMPT.format(
        resume_text=resume_text,
        job_description=description,
        company=company,
        role=role,
        missing_keywords=missing_keywords or "none listed",
    )

    logger.info("Generating tailored resume for %s @ %s", role, company)
    raw = _call_ai(prompt)
    if not raw:
        logger.error("AI returned no content for resume tailoring")
        return None

    # Parse JSON — handle markdown code fences
    json_str = raw.strip()
    if json_str.startswith("```"):
        lines = json_str.splitlines()
        json_str = "\n".join(lines[1:-1]) if lines[-1].strip() == "```" else "\n".join(lines[1:])

    try:
        data = json.loads(json_str)
    except json.JSONDecodeError as e:
        logger.error("Failed to parse AI JSON: %s\nRaw: %s", e, raw[:500])
        return None

    doc = _build_docx(data)
    if doc is None:
        return None

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_company = "".join(c for c in company if c.isalnum() or c in "_-")[:20]
    safe_role = "".join(c for c in role if c.isalnum() or c in "_-")[:30]
    archive_dir = output_dir / "archive"
    archive_dir.mkdir(parents=True, exist_ok=True)
    filename = f"resume_{safe_company}_{safe_role}_{ts}.docx"
    out_path = archive_dir / filename

    try:
        doc.save(str(out_path))
        logger.info("Saved tailored resume: %s", out_path)
        return out_path
    except Exception as e:
        logger.error("Failed to save DOCX: %s", e)
        return None
