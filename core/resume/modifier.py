"""Modify DOCX resumes with tailored content."""

from typing import Optional
from datetime import datetime


class ResumeModifier:
    """Modify DOCX resumes programmatically."""

    def create_tailored_version(
        self,
        original_path: str,
        tailored_text: str,
        job_company: str,
        output_dir: str = "data"
    ) -> Optional[str]:
        """
        Create a tailored resume DOCX from tailored text.

        Args:
            original_path: Path to original resume DOCX
            tailored_text: Tailored resume text from AI
            job_company: Company name for filename
            output_dir: Directory to save tailored resume

        Returns:
            Path to saved tailored resume, or None if failed
        """
        try:
            from docx import Document

            # Create new document from tailored text
            doc = Document()

            # Split text into paragraphs and add to document
            for paragraph_text in tailored_text.split("\n"):
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text)

            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"resume_tailored_{job_company}_{timestamp}.docx"
            filepath = f"{output_dir}/{filename}"

            # Create output directory if needed
            import os
            os.makedirs(output_dir, exist_ok=True)

            doc.save(filepath)
            return filepath
        except Exception as e:
            print(f"Error creating tailored resume: {e}")
            return None

    def update_section(
        self,
        resume_path: str,
        section_name: str,
        new_content: str
    ) -> bool:
        """Update a specific section in resume."""
        try:
            from docx import Document
            doc = Document(resume_path)

            # Find and replace section
            in_section = False
            section_start = None

            for i, para in enumerate(doc.paragraphs):
                if section_name.lower() in para.text.lower():
                    in_section = True
                    section_start = i
                elif in_section and para.text.strip() and para.text[0].isupper():
                    # New section found
                    break

            if section_start is not None:
                # Remove old content in section
                for i in range(len(doc.paragraphs) - 1, section_start, -1):
                    if section_start < i < len(doc.paragraphs):
                        p = doc.paragraphs[i]._element
                        p.getparent().remove(p)

                # Add new content
                for line in new_content.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)

                doc.save(resume_path)
                return True
        except Exception as e:
            print(f"Error updating resume section: {e}")

        return False
