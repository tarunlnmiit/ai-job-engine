"""Parse resume files to extract text and metadata."""

from typing import Optional
import os
from logger import get_logger

logger = get_logger("resume.parser")


class ResumeParser:
    """Parse resume files (DOCX, PDF, TXT)."""

    def parse(self, filepath: str) -> Optional[str]:
        """Parse resume file and return text content."""
        logger.info("Parsing resume: %s", filepath)
        if not os.path.exists(filepath):
            logger.error("Resume file not found: %s", filepath)
            return None

        ext = os.path.splitext(filepath)[1].lower()
        logger.debug("Resume format: %s", ext)

        if ext == ".docx":
            return self._parse_docx(filepath)
        elif ext == ".pdf":
            return self._parse_pdf(filepath)
        elif ext == ".txt":
            return self._parse_txt(filepath)
        else:
            logger.error("Unsupported resume format: %s", ext)
            return None

    def _parse_docx(self, filepath: str) -> Optional[str]:
        """Parse DOCX file."""
        try:
            from docx import Document
            doc = Document(filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
            logger.info("DOCX parsed — %d chars, %d paragraphs", len(text), len(doc.paragraphs))
            return text
        except Exception as e:
            logger.error("Error parsing DOCX: %s", e)
            return None

    def _parse_pdf(self, filepath: str) -> Optional[str]:
        """Parse PDF file."""
        try:
            import pdfplumber
            text = ""
            with pdfplumber.open(filepath) as pdf:
                logger.debug("PDF has %d pages", len(pdf.pages))
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    logger.debug("Page %d: %d chars", i + 1, len(page_text))
                    text += page_text + "\n"
            logger.info("PDF parsed — %d chars total", len(text))
            return text
        except Exception as e:
            logger.error("Error parsing PDF: %s", e)
            return None

    def _parse_txt(self, filepath: str) -> str:
        """Parse TXT file."""
        with open(filepath, "r") as f:
            text = f.read()
        logger.info("TXT parsed — %d chars", len(text))
        return text
